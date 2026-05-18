"""Build the elevated EMERALD-AI dissertation proposal as a Word document.

v0.3 (2026-05-18) — third draft. Rebuilt against the post-bot literature brain
(108 papers, 1087 authors). Adds the climate–credit risk channel, the
sociotechnical fairness critique, the reject-inference / selection-bias
treatment, MLOps deployment evidence, and the opacity-of-finance framing
that v0.2 left implicit. Target length ~8,000 words.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x2E)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t


doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ============================================================
# 1. TITLE BLOCK
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(
    "EMERALD-AI: An Explainable, Calibrated, and Audit-Ready\n"
    "Machine Learning Framework for Green Loan Credit Scoring,\n"
    "Operationalised as a Full-Stack Decision-Support Platform"
)
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x14, 0x4F, 0x32)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("MSc Applied Artificial Intelligence — Dissertation Proposal (Third Draft)")
sr.italic = True
sr.font.size = Pt(11)

doc.add_paragraph()

# ============================================================
# 2. ABSTRACT
# ============================================================
add_heading(doc, "Abstract", 1)
add_para(
    doc,
    "The global green loan market has grown at a compound annual rate exceeding 40% since 2018, "
    "yet credit-risk infrastructure for sustainable lending remains anchored in scorecards designed "
    "for conventional consumer credit. This is no longer only a sustainability concern: recent "
    "macro-financial evidence shows that climate change and clean-technology capital expenditure now "
    "enter credit-risk cycles in measurable, asset-class-specific ways [Mirza et al., 2024], and that "
    "regulatory climate-risk disclosures move credit-default-swap term structure [Kölbel et al., 2022]. "
    "This dissertation proposes EMERALD-AI, an explainable, calibrated, and audit-ready machine "
    "learning (ML) framework for green-loan credit scoring, evaluated on a proprietary 2019 dataset of "
    "14,135 funded transactions with 166 structured variables. The framework benchmarks a coherent "
    "family of classifiers spanning regularised logistic regression, kernel SVM, gradient-boosted "
    "decision trees (XGBoost, LightGBM, CatBoost), and modern tabular deep learning (FT-Transformer, "
    "MLP) under nested stratified cross-validation with Bayesian hyperparameter optimisation. To meet "
    "EU AI Act and FCA Consumer Duty obligations, EMERALD-AI integrates (i) a multi-layer "
    "explainability stack (TreeSHAP, KernelSHAP, counterfactual DiCE), (ii) Platt and isotonic "
    "probability calibration, (iii) split-conformal prediction for distribution-free uncertainty "
    "quantification, (iv) a fairness audit covering disparate impact, equalised odds, and calibration "
    "parity across industry and geographic proxies, disciplined by Selbst et al.'s (2019) "
    "sociotechnical critique of purely-technical fairness interventions, and (v) an explicit reject-"
    "inference treatment of accepted-only selection bias [Banasik and Crook, 2007; Kang et al., 2021; "
    "Shen et al., 2022]. The trained system is operationalised through a containerised FastAPI/React "
    "full-stack application supporting single-applicant scoring with local explanations, batch CSV "
    "scoring, portfolio analytics, and a SHAP-based feature explorer, with the deployment architecture "
    "informed by Paleyes et al.'s (2022) survey of production-ML failure modes. Expected "
    "contributions: (a) the first published, empirically validated, green-loan-specific ML credit-"
    "scoring framework; (b) a reproducible MLOps blueprint for high-stakes financial AI; (c) a "
    "regulator-ready audit template covering calibration, fairness, robustness, and explanation "
    "fidelity; (d) an open-source platform that bridges research prototypes and lending-officer "
    "workflows."
)

add_para(doc, "Alternative working titles under consideration:", italic=True)
add_bullet(doc, "VERDIS-AI — Verifiable Explainable Risk Decisioning for Industrial Sustainability")
add_bullet(doc, "SYLVAN-Credit — Sustainable, Yield-aware Lending with Verifiable, Auditable Networks")
add_bullet(doc, "TERRA-Score — Transparent, Equitable, Regulator-aligned Risk Assessment for green lending")

doc.add_page_break()

# ============================================================
# 3. INTRODUCTION
# ============================================================
add_heading(doc, "1. Introduction", 1)
add_para(
    doc,
    "Green lending has expanded from a niche sustainability initiative into a mainstream fixture of "
    "modern finance. Green loans are credit products whose proceeds are ring-fenced for environmentally "
    "beneficial purposes — renewable energy installations, energy-efficient building upgrades, clean "
    "transportation, and sustainable agriculture, among others. The Loan Market Association et al. "
    "(2023) reports that the global green loan market has grown at a compound annual rate exceeding "
    "40% since 2018, reflecting both rising investor appetite for responsible assets and intensifying "
    "regulatory pressure on financial institutions to align their portfolios with net-zero commitments."
)
add_para(
    doc,
    "Despite this rapid growth, credit assessment in green lending remains anchored to scoring "
    "frameworks built for conventional loan products. Traditional models rely heavily on a narrow set "
    "of indicators — repayment history, income, and debt coverage — and are poorly equipped to "
    "evaluate the environmental viability of funded projects, the long-horizon cash-flow profiles of "
    "renewable-energy assets, or the reputational exposure arising from greenwashing. As Ehlers and "
    "Packer (2017) observe, the information gap between borrowers and lenders in green debt markets "
    "is one of the defining structural weaknesses of the asset class, creating risks of capital "
    "misallocation that undermine the very purpose of sustainable finance. The mis-pricing argument is "
    "no longer purely theoretical: Mirza et al. (2024) show, on a panel of European Union credit-risk "
    "indicators, that climate change and clean-technology capital expenditure measurably alter "
    "credit-risk cycles, while Kölbel et al. (2022) demonstrate using BERT-based text analysis that "
    "regulatory climate-risk disclosures move the term structure of credit-default swaps. The climate-"
    "credit channel that green-loan portfolios sit on is now empirically observable in market prices, "
    "and lenders that fail to model it explicitly carry an unquantified concentration of climate-"
    "transition risk."
)
add_para(
    doc,
    "Artificial Intelligence (AI) and Machine Learning offer a fundamentally different approach. "
    "Rather than applying fixed rules to a small set of variables, ML algorithms can simultaneously "
    "process hundreds of features, identify non-linear relationships that conventional models miss, "
    "and produce risk assessments that adapt as new data becomes available. This dissertation presents "
    "EMERALD-AI — an explainable, calibrated, and audit-ready ML credit-scoring framework trained on a "
    "real-world dataset of 14,135 funded green-loan transactions from 2019, encompassing 166 "
    "structured variables including credit scores, loan amounts, annual percentage rates, factor "
    "rates, payback ratios, industry classifications, borrower states, deal stages, and tier "
    "classifications. Crucially, the research also delivers a production-ready full-stack web "
    "application through which the trained model can be accessed interactively, making the findings "
    "directly usable by lending institutions rather than confined to academic outputs."
)

# ============================================================
# 4. MOTIVATION
# ============================================================
add_heading(doc, "2. Motivation", 1)

add_heading(doc, "2.1 Operational Gap in Green Lending Practice", 2)
add_para(
    doc,
    "Preliminary inspection of the 2019 dataset reveals striking heterogeneity across borrowers: "
    "credit scores span the full sub-prime to super-prime range, APR structures vary materially by "
    "industry and tier, and deal outcomes are unevenly distributed across states and sectors. Despite "
    "this heterogeneity, green-lending institutions commonly apply broadly uniform credit thresholds. "
    "This is the canonical inefficiency that ML is designed to address — learning to treat different "
    "borrower profiles differently in ways that are grounded in data rather than intuition."
)

add_heading(doc, "2.2 Gap in the Academic Literature", 2)
add_para(
    doc,
    "There is a well-established body of work on ML in conventional credit scoring [Baesens et al., "
    "2003; Lessmann et al., 2015; Dumitrescu et al., 2022], and a growing literature on green finance "
    "[Flammer, 2021; Tang and Zhang, 2020; Gan et al., 2023]. There is also a recent and rapidly "
    "maturing literature on the climate–credit-risk channel itself [Kölbel et al., 2022; Mirza et al., "
    "2024; Dafermos and Nikolaidi, 2021]. However, empirical research at the intersection of these "
    "three fields — work that specifically asks how modern ML can improve credit decisions in green "
    "lending, on real green-loan data, with the calibration, explainability, and fairness machinery "
    "expected of high-risk AI under the EU AI Act — is conspicuously absent. EMERALD-AI directly "
    "addresses this triple intersection."
)

add_heading(doc, "2.3 Regulatory Pressure for Explainability and Audit", 2)
add_para(
    doc,
    "A model that produces accurate predictions but cannot explain them is not deployable in "
    "regulated financial services. The EU AI Act [European Commission, 2021] classifies credit "
    "scoring as an Annex III high-risk application, mandating data governance, technical "
    "documentation, transparency, human oversight, accuracy and robustness logging, and post-market "
    "monitoring. The UK Financial Conduct Authority's Consumer Duty [FCA, 2022] requires firms to "
    "demonstrate that automated decisions are fair and explainable to consumers. The EBA Guidelines "
    "on Loan Origination and Monitoring [EBA, 2020] further require lenders to evidence model "
    "robustness and interpretability. The pressure here is not solely technical: Pasquale (2015) and "
    "Burrell (2016) argue persuasively that algorithmic opacity in finance is a sociotechnical "
    "problem in its own right, eroding the accountability assumptions on which retail credit markets "
    "depend. EMERALD-AI treats explainability, calibration, and fairness as non-negotiable first-"
    "class design requirements — embedded from the outset rather than retrofitted."
)

add_heading(doc, "2.4 Contribution to Climate Action", 2)
add_para(
    doc,
    "Poorly allocated green finance is not just a financial problem — it is an environmental one. "
    "Capital directed toward projects that fail to deliver genuine ecological benefit delays the "
    "adoption of clean technologies and undermines public trust in sustainable finance. Improving "
    "green credit decisions, even incrementally, contributes to the United Nations Sustainable "
    "Development Goals — particularly SDG 7 (Affordable and Clean Energy), SDG 11 (Sustainable "
    "Cities and Communities), and SDG 13 (Climate Action) — and supports the financial system's "
    "broader alignment with the Paris Agreement. Macroprudential proposals to introduce green-"
    "differentiated capital requirements [Dafermos and Nikolaidi, 2021] further suggest that the "
    "regulatory cost of holding green vs. brown exposures may diverge in coming years, so an "
    "auditable, green-specific risk model is also a forward-looking operational hedge."
)

# ============================================================
# 5. RESEARCH OBJECTIVES AND QUESTIONS
# ============================================================
add_heading(doc, "3. Research Objectives and Research Questions", 1)

add_heading(doc, "3.1 Research Objectives", 2)
add_numbered(doc, "Conduct a rigorous exploratory analysis of the 2019 All Funded Green Loan dataset, including a systematic target-leakage audit and an accepted-only selection-bias diagnostic, to identify the statistical distributions, inter-feature relationships, and pre-funding variables most strongly associated with green-loan creditworthiness.")
add_numbered(doc, "Design and implement a fully reproducible, version-controlled end-to-end ML pipeline that integrates financial, operational, and borrower-level features into a credit-scoring model purpose-built for green-loan assessment.")
add_numbered(doc, "Train, optimise, and benchmark six ML algorithms — L1/L2-regularised Logistic Regression, kernel SVM, Random Forest, XGBoost, LightGBM/CatBoost, and a tabular deep-learning family (MLP and FT-Transformer) — under nested stratified cross-validation with Bayesian hyperparameter optimisation, comparing them using credit-risk-appropriate metrics with statistical significance testing.")
add_numbered(doc, "Integrate a multi-layer explainability stack (TreeSHAP, KernelSHAP, LIME, DiCE counterfactuals) and empirically validate explanation fidelity to produce transparent decisions at both global and local scales, in compliance with EU AI Act and FCA Consumer Duty requirements.")
add_numbered(doc, "Apply post-hoc probability calibration (Platt, isotonic, temperature scaling) and split-conformal prediction to deliver well-calibrated probabilities and distribution-free per-applicant uncertainty intervals — properties largely absent in the existing green-credit ML literature.")
add_numbered(doc, "Conduct a fairness, stability, and robustness audit across industry, geographic, and business-size proxy axes, measuring disparate impact, equalised-odds gaps, predictive parity, calibration-within-group, and prediction stability under input perturbation and concept drift, with the audit framed against Selbst et al.'s (2019) five sociotechnical 'traps' rather than as a purely technical exercise.")
add_numbered(doc, "Develop and deploy a containerised, MLOps-grade full-stack web application — the EMERALD-AI platform — that makes the trained model accessible through a professional interactive interface, including single-borrower prediction with local explanation and counterfactual, batch CSV scoring, portfolio analytics with fairness panel, and a SHAP feature explorer.")

add_heading(doc, "3.2 Research Questions", 2)
add_bullet(doc, "RQ1 — Which pre-funding variables in the 2019 green-loan dataset are the strongest predictors of borrower creditworthiness, and do they differ meaningfully from the predictors that dominate conventional credit scoring?")
add_bullet(doc, "RQ2 — How do regularised linear, tree-ensemble, and modern tabular deep-learning architectures compare on green-loan data when measured on PR-AUC, ROC-AUC, KS, Brier score, Expected Calibration Error, and recall@top-decile, and how do those differences persist after post-hoc calibration?")
add_bullet(doc, "RQ3 — How robust and fair are the candidate models across industry, geography, and business-size segments, and to what extent does the chosen explainability stack produce attributions that are faithful, stable, and counterfactually actionable?")
add_bullet(doc, "RQ4 — Does the EMERALD-AI web application effectively communicate AI credit-scoring outputs — predictions, conformal intervals, SHAP attributions, and counterfactual recourses — to non-technical lending stakeholders in a manner consistent with EU AI Act Article 13 transparency obligations?")

doc.add_page_break()

# ============================================================
# 6. LITERATURE REVIEW (EXPANDED)
# ============================================================
add_heading(doc, "4. Literature Review", 1)

add_heading(doc, "4.1 The Evolution of Credit Scoring: From Discriminant Analysis to Statistical Learning", 2)
add_para(
    doc,
    "Modern credit scoring traces its lineage to Altman's Z-score [Altman, 1968] and Beaver's "
    "univariate ratio analyses [Beaver, 1966], later formalised through linear discriminant analysis "
    "and, eventually, logistic regression — which has remained the regulatory default for over four "
    "decades owing to its monotonicity, sparsity, and direct mapping to odds and points-based "
    "scorecards [Hand and Henley, 1997; Henley, 1995; Thomas, 2000]. The migration to non-parametric "
    "learning began with neural networks and SVMs in the late 1990s, but adoption stalled on opacity "
    "grounds: regulators viewed black-box models as incompatible with the right-to-explanation "
    "regimes embedded in ECOA, GDPR Article 22, and Basel risk-weight reporting. Baesens et al. "
    "(2003) carried out one of the first systematic comparative evaluations of classification "
    "algorithms on real credit datasets, demonstrating that neural networks and SVMs consistently "
    "outperformed logistic regression. Lessmann et al. (2015) extended this benchmark to 41 "
    "classifiers across eight credit datasets and established two findings that still shape the "
    "field: heterogeneous ensemble methods (random forest, gradient boosting, stacking) dominate "
    "single learners, and the marginal accuracy gains from architectural sophistication shrink as "
    "datasets grow. Dumitrescu et al. (2022) recently demonstrated that interpretable logistic "
    "regression augmented with explicit tree-derived interaction terms can close most of the gap to "
    "XGBoost — reframing the explainability/accuracy trade-off as a design choice rather than a "
    "fixed frontier. Industry adoption, however, has lagged the empirical evidence by a margin that "
    "no purely-technical argument can explain: vendor lock-in, model-validation cost, and the cost "
    "of re-papering existing scorecards continue to defer the migration even where the methodology "
    "case is unambiguous."
)

add_heading(doc, "4.2 Gradient-Boosted Decision Trees as the Tabular State of the Art", 2)
add_para(
    doc,
    "For heterogeneous tabular data with mixed-type features, missing values, and moderate scale "
    "(10³–10⁶ rows), gradient-boosted decision trees (GBDT) have been the consistent empirical "
    "winner. XGBoost [Chen and Guestrin, 2016] introduced second-order Newton boosting with L1/L2 "
    "regularisation and sparsity-aware split finding; LightGBM [Ke et al., 2017] added "
    "histogram-based binning and leaf-wise growth for order-of-magnitude training speedups; CatBoost "
    "[Prokhorenkova et al., 2018] introduced ordered boosting and target-statistics encoding that "
    "eliminate the prediction-shift bias arising from naive target encoding. Independent benchmarks "
    "consistently rank GBDTs at parity with — and often above — deep tabular architectures on "
    "classification problems of the scale considered here [Shwartz-Ziv and Armon, 2022; Grinsztajn "
    "et al., 2022]. Two properties make GBDTs particularly well-suited to credit risk: (i) native "
    "support for monotonic constraints, which let modellers enforce the directionality of features "
    "(e.g., higher credit score must not decrease creditworthiness probability) that supervisors "
    "have been signalling for over a decade through model-risk-management guidance such as Fed SR "
    "11-7 and EBA GL/2017/16; and (ii) integration with TreeSHAP [Lundberg et al., 2020], which "
    "yields exact Shapley attributions in polynomial time — making per-decision explanations cheap "
    "enough to ship in production."
)

add_heading(doc, "4.3 Tabular Deep Learning: Promise, Pitfalls, and Recent Advances", 2)
add_para(
    doc,
    "The push to extend deep learning's vision and language successes to tabular data has produced a "
    "rapidly evolving family of architectures. TabNet [Arik and Pfister, 2021] introduced sequential "
    "attention with sparse feature masks that double as built-in feature attribution; FT-Transformer "
    "[Gorishniy et al., 2021] adapted the Transformer encoder to tabular features via "
    "feature-tokenisation; SAINT [Somepalli et al., 2021] combined row- and column-attention with "
    "contrastive pre-training; NODE [Popov et al., 2020] built differentiable oblivious decision-tree "
    "ensembles trainable end-to-end. The empirical picture, however, is more sober: on most public "
    "benchmarks of moderate size these architectures fail to beat well-tuned GBDTs [Borisov et al., "
    "2022; Grinsztajn et al., 2022]. Deep tabular models earn their keep in specific regimes — very "
    "large datasets, the presence of high-cardinality categoricals where learned embeddings outperform "
    "target encoding, and multi-task or self-supervised pre-training scenarios where representation "
    "transfer matters. For green credit scoring at the scale of approximately 14,000 transactions, "
    "this dissertation includes a tuned FT-Transformer and MLP as honest comparators while resisting "
    "the temptation to over-claim deep learning's necessity. Tabular foundation models such as "
    "TabPFN remain a frontier development worth tracking but are not yet directly applicable: their "
    "current ≤1,000-sample constraint sits an order of magnitude below the green-loan dataset size."
)

add_heading(doc, "4.4 Class Imbalance, Calibration, Uncertainty, and Reject Inference in Credit Risk", 2)
add_para(
    doc,
    "Credit-default datasets are intrinsically imbalanced — typical positive (default) prevalence in "
    "well-underwritten portfolios sits in the 2–15% range. Naive training on imbalanced data produces "
    "classifiers that minimise overall error but underperform on the minority class that drives loss. "
    "Standard interventions include resampling (SMOTE [Chawla et al., 2002], SMOTE-NC for mixed types, "
    "ADASYN [He et al., 2008]), cost-sensitive loss reweighting, and focal loss [Lin et al., 2017]. "
    "Critically, resampling must occur strictly inside cross-validation folds to prevent test-set "
    "contamination — a methodological mistake that is widespread in published credit-scoring work "
    "[Santos et al., 2018]. A separate and often-overlooked concern is calibration: a model can rank "
    "applicants correctly (high AUC) while producing systematically biased probability estimates, "
    "which breaks expected-loss pricing, capital provisioning, and regulatory probability-of-default "
    "reporting under the Internal Ratings-Based approach. Platt scaling [Platt, 1999], isotonic "
    "regression [Zadrozny and Elkan, 2002], temperature scaling, and conformal prediction [Vovk et "
    "al., 2005; Angelopoulos and Bates, 2023] provide post-hoc calibration and distribution-free "
    "uncertainty quantification respectively, and are largely absent from existing green-credit "
    "literature. EMERALD-AI treats both as mandatory pipeline stages."
)
add_para(
    doc,
    "A third and historically under-treated problem — accepted-only selection bias, or 'reject "
    "inference' — sits adjacent to the imbalance question. The training distribution observed by any "
    "credit scoring model is drawn only from applicants the prior underwriting policy accepted; "
    "applicants who were declined never generate the outcome labels needed for supervised learning, "
    "so the model is fitted on a non-random sub-population whose covariates correlate with the very "
    "decisions the model is replacing. Banasik and Crook (2007) provide the canonical operational-"
    "research treatment, evaluating augmentation, parcelling, and Heckman-style sample-selection "
    "corrections on UK credit data. The 2020s literature has substantially modernised the toolkit: "
    "Kang et al. (2021) propose a graph-based semi-supervised framework that propagates labels from "
    "accepted to rejected applicants while explicitly handling imbalance, and Shen et al. (2022) "
    "combine three-way decisioning with a safe semi-supervised support vector machine to reduce "
    "label-propagation error. EMERALD-AI does not claim to resolve the reject-inference problem "
    "(the dataset contains only funded transactions and therefore lacks the unfunded counterfactual), "
    "but it surfaces the limitation explicitly, follows Banasik and Crook's sensitivity-analysis "
    "protocol, and discusses what an extended pipeline would require — as detailed in §5.2."
)

add_heading(doc, "4.5 Green Finance, ESG Risk, and the Specificity of Sustainable Lending", 2)
add_para(
    doc,
    "Sustainable lending differs from conventional consumer or commercial credit along three "
    "structural dimensions that current scorecards under-model. First, the cash-flow profile of "
    "green-purposed assets — solar PV, building retrofits, EV fleets — is back-loaded and "
    "policy-sensitive, exposing lenders to long-horizon regulatory and technology risk that linear "
    "repayment history does not capture [Flammer, 2021; Tang and Zhang, 2020]. Second, green loans "
    "carry a non-financial reputational tail: greenwashing exposure can crystallise into both "
    "regulatory penalty (EU Taxonomy disclosure, SFDR) and franchise damage that materially affects "
    "loss-given-default outside the borrower's own balance sheet [Ehlers and Packer, 2017]. Third, "
    "what counts as 'green' is heterogeneous across jurisdictions and evolving — the EU Taxonomy "
    "[European Commission, 2020], the Climate Bonds Standard [CBI, 2023], and emerging national "
    "taxonomies each define eligibility differently, introducing label noise into any portfolio "
    "assembled across regulatory regimes. ESG scores themselves exhibit substantial inter-rater "
    "disagreement [Berg et al., 2022], making them unreliable as direct features. EMERALD-AI does "
    "not use third-party ESG ratings as model inputs for this reason."
)
add_para(
    doc,
    "Crucially, the green-versus-conventional distinction is no longer purely a definitional or "
    "reputational matter: it is becoming empirically observable in credit-risk indicators. Mirza et "
    "al. (2024) show on European Union panel data that climate-change exposure and clean-technology "
    "capital expenditure measurably alter credit-risk cycles, while Kölbel et al. (2022) use BERT to "
    "quantify regulatory climate-risk disclosures and demonstrate that they move the term structure "
    "of credit-default swaps. On the borrower side, Gan et al. (2023) document that green-bond "
    "issuance in the Chinese market expands trade-credit access for issuers — evidence that the "
    "green label itself carries credit-relevant information. Dafermos and Nikolaidi (2021) develop a "
    "macrofinancial model in which green-differentiated capital requirements materially reduce "
    "system-wide climate risk, suggesting that supervisors may, within the dissertation's deployment "
    "horizon, begin to require lenders to separate green from non-green exposures in their internal "
    "models. To date no published work has built a credit-scoring framework that engages with all "
    "three structural specifics of green lending and the climate–credit channel on real green-loan "
    "transaction data."
)

add_heading(doc, "4.6 Explainable AI in High-Stakes Finance", 2)
add_para(
    doc,
    "The case for explainability in financial ML is sociotechnical before it is technical. Burrell "
    "(2016) decomposes algorithmic opacity into three distinct sources — intentional commercial "
    "secrecy, technical illiteracy among affected users, and the irreducible scale-and-complexity of "
    "high-dimensional ML — and shows that interventions targeting one source leave the others "
    "untouched. Pasquale (2015) frames the political-economy stakes: a financial system that scores, "
    "ranks, and prices people through opaque algorithms erodes the accountability assumptions on "
    "which consumer trust and supervisory review depend. The XAI stack below is the technical answer "
    "to Burrell's third opacity source; the regulatory pressure for explainability is driven by "
    "all three."
)
add_para(
    doc,
    "Explainability in ML credit scoring is not a single technique but a stack of complementary "
    "tools, each addressing a different audit question. Global feature attribution (permutation "
    "importance, mean |SHAP|) answers 'what does the model rely on at the portfolio level?'; local "
    "attribution (LIME [Ribeiro et al., 2016], TreeSHAP) answers 'why did this borrower receive this "
    "score?'; counterfactual explanations (DiCE [Mothilal et al., 2020], Wachter-style counterfactuals "
    "[Wachter et al., 2017]) answer 'what would the borrower need to change to be approved?' — the "
    "form regulators and consumer-advocacy bodies increasingly favour because it is actionable. "
    "SHAP's appeal derives from its grounding in cooperative game theory: it is the unique attribution "
    "method satisfying local accuracy, missingness, and consistency [Lundberg and Lee, 2017]. Critical "
    "perspectives, however, are growing. Feature attribution can be misleading under feature "
    "correlation [Aas et al., 2021; Janzing et al., 2020]; counterfactuals can be unstable [Rawal and "
    "Lakkaraju, 2020]; and post-hoc explanations of opaque models can be deliberately manipulated to "
    "hide bias [Slack et al., 2020]. Rudin (2019) goes further and argues that for high-stakes "
    "decisions, inherently interpretable models (constrained additive models, sparse decision lists) "
    "should be preferred over post-hoc explanations of opaque ones. EMERALD-AI takes the pragmatic "
    "middle path: it deploys a constrained tree ensemble whose structure is itself inspectable, and "
    "layers a multi-method explanation stack (TreeSHAP + KernelSHAP cross-check + counterfactuals + "
    "monotonic-constraint verification) rather than relying on a single attribution mechanism. "
    "Explanation faithfulness is itself empirically validated via the Quantus toolkit [Hedström et "
    "al., 2023]. The audit artefacts themselves — model card, datasheet, and method card — follow "
    "the prescriptive transparency template proposed by Adkins et al. (2022), which extends model "
    "cards [Mitchell et al., 2019] and datasheets [Gebru et al., 2021] with explicit guidance on "
    "what *not* to do with the model rather than only how it was built."
)

add_heading(doc, "4.7 Fairness, Robustness, and Regulatory Compliance", 2)
add_para(
    doc,
    "Fairness in credit scoring is operationalised through statistical parity, equalised odds [Hardt "
    "et al., 2016], predictive parity, and calibration within groups — definitions that are provably "
    "incompatible in non-trivial cases [Chouldechova, 2017; Kleinberg et al., 2017], forcing "
    "modellers to choose which parity criterion best fits the deployment context. In the green "
    "lending case, protected attributes (race, gender) are not directly available, but proxies "
    "(industry, geography, business size) can produce indirect disparate impact testable under ECOA "
    "and equivalent UK regulation. The EU AI Act [European Commission, 2021] classifies credit "
    "scoring as Annex III high-risk, triggering obligations around data governance, technical "
    "documentation, human oversight, accuracy and robustness logging, and post-market monitoring; "
    "the FCA Consumer Duty [FCA, 2022] adds outcome-based fairness obligations; and the EBA "
    "Guidelines on Loan Origination and Monitoring [EBA, 2020] require lenders to demonstrate model "
    "robustness and explainability. Robustness extends beyond fairness to stability under input "
    "perturbation, adversarial input resistance, and concept-drift monitoring [Lu et al., 2018]."
)
add_para(
    doc,
    "A purely-technical fairness audit, however well-instrumented, is not enough. Selbst et al. "
    "(2019) identify five abstraction 'traps' that defeat fair-ML interventions when they ignore the "
    "social context the model is embedded in: the framing trap (treating the technical system as the "
    "entire decision pipeline), the portability trap (porting fairness fixes across deployment "
    "contexts without re-validation), the formalism trap (assuming a single mathematical fairness "
    "criterion captures a contested social concept), the ripple-effect trap (ignoring how the model "
    "reshapes the system around it), and the solutionism trap (assuming an ML system is the right "
    "tool for the problem at all). EMERALD-AI's audit therefore reports the chosen parity criterion "
    "as a documented value judgement, surfaces the underlying base-rate decomposition so reviewers "
    "can disagree on grounds the audit makes visible, and frames its fairness claims as conditional "
    "on the deployment context rather than as portable guarantees. This audit framework is "
    "integrated into the development loop rather than treated as a post-hoc checklist."
)

add_heading(doc, "4.8 Identified Research Gap and the Positioning of EMERALD-AI", 2)
add_para(
    doc,
    "Synthesising the bodies of literature above, the gap is precise: there is no published, "
    "empirically validated, regulator-ready ML credit-scoring framework trained on real green-loan "
    "transaction data that simultaneously integrates (i) modern tabular learners benchmarked against "
    "interpretable baselines under identical preprocessing, (ii) post-hoc calibration and "
    "distribution-free uncertainty quantification, (iii) a multi-method explainability stack with "
    "empirical fidelity validation, (iv) a sociotechnically-disciplined fairness and robustness audit "
    "appropriate to green lending's industry and geography proxies, (v) an explicit acknowledgement "
    "and partial treatment of accepted-only selection bias, and (vi) a deployable interface that "
    "lets lending officers act on outputs. EMERALD-AI is positioned to occupy this gap end-to-end. "
    "The 'first' claim is defended against a documented OpenAlex-mediated systematic search "
    "(retained in the project's gaps log) which surfaced adjacent but non-overlapping work on "
    "conventional consumer credit ML, on green-bond pricing, and on the climate-credit channel — "
    "but no work combining all six requirements on green-loan data."
)

doc.add_page_break()

# ============================================================
# 7. METHODOLOGY (EXPANDED)
# ============================================================
add_heading(doc, "5. Methodology", 1)

add_heading(doc, "5.1 Research Design and Epistemological Position", 2)
add_para(
    doc,
    "EMERALD-AI is a quantitative, post-positivist applied research project structured around the "
    "CRISP-DM lifecycle [Wirth and Hipp, 2000] and instantiated as a reproducible, version-controlled "
    "MLOps pipeline. It progresses through business understanding (Sections 1–2), data understanding "
    "(5.3–5.4), data preparation (5.5–5.7), modelling (5.8–5.9), evaluation (5.10–5.13), and "
    "deployment (5.14), with a feedback loop from deployment back to modelling enabling iterative "
    "refinement. The unit of analysis is the individual funded green-loan transaction; the inferential "
    "target is the conditional probability of creditworthiness given pre-funding observable features."
)

add_heading(doc, "5.2 Problem Formalisation, Label Construction, and Selection Bias", 2)
add_para(
    doc,
    "Let X ∈ ℝ^d denote the feature vector of a green-loan applicant and Y ∈ {0,1} the binary "
    "creditworthiness indicator, with Y=1 corresponding to favourable performance (paidOff ∪ current) "
    "and Y=0 corresponding to delinquency (default ∪ behind). The estimand is the conditional "
    "probability η(x) = ℙ(Y=1 | X=x), approximated by the learned model f̂(x). The Deal Status field "
    "provides direct supervision for 14,022 of the 14,135 records (99.2% labelled coverage). Three "
    "label-construction risks are explicitly addressed:"
)
add_bullet(doc, "Censoring bias. Current loans are right-censored in their observation window and may yet default, biasing the Y=1 class. Mitigation: a sensitivity analysis that excludes still-active deals from the training distribution and reports performance with and without their inclusion.")
add_bullet(doc, "Definition leakage. Fields populated only after the deal outcome is known (terminal Deal Stage, post-funding adjustment fields) must be excluded from X; the leakage audit (§5.3) catalogues every such field.")
add_bullet(doc, "Accepted-only selection bias (reject inference). The dataset contains only funded loans — applicants the prior underwriting policy declined are absent, so the empirical distribution of (X, Y) is conditional on prior acceptance. Banasik and Crook (2007) showed on UK credit data that this can materially bias parameter estimates if ignored. The contemporary toolkit — graph-based semi-supervised label propagation [Kang et al., 2021] and three-way-decision safe semi-supervised SVMs [Shen et al., 2022] — assumes access to the rejected applicants' features, which this dataset does not provide. EMERALD-AI therefore (a) reports the limitation transparently in the dissertation's external-validity section, (b) executes Banasik and Crook's sensitivity-analysis protocol of comparing tight and loose acceptance-window cohorts where dataset metadata supports the cut, and (c) documents the extended pipeline — joint observation of accepted and declined applicants — as a stipulated requirement of any production deployment.")
add_para(
    doc,
    "A coherent set of related questions is noted but explicitly scoped out, each being a dissertation "
    "in its own right: federated and decentralised credit scoring across institutions, differential-"
    "privacy guarantees on per-applicant attributions, secure multi-party computation for cross-"
    "lender risk pooling, survival-analytic time-to-default modelling, and online / continual "
    "learning under regulatory-grade post-market monitoring. These directions are recorded in §8 as "
    "future work rather than addressed here."
)

add_heading(doc, "5.3 Dataset Characterisation and Target-Leakage Audit", 2)
add_para(
    doc,
    "The 2019 All Funded Green Loan dataset contains 14,135 transactions across 166 columns. Each "
    "column is classified into one of six categories: (a) pre-funding applicant attributes (Credit "
    "Score, Annual Revenue, Time in Business, Average Monthly Sales, Industry, Borrower State); (b) "
    "pre-funding loan-offer attributes (Amount Sought, Amount Funded, APR, Factor Rate, Term, "
    "Payback, Commission); (c) loan structural metadata (Lender, Product Type, Current Tier); (d) "
    "deal-progression timestamps (Start, Offer Received, Contract Signed, Deal Closed); (e) "
    "post-funding observed outcomes (Deal Status, Payments Made, Days Past Due — these define Y and "
    "must not enter X); (f) administrative or free-text fields excluded as features. Only categories "
    "(a)–(d) enter the feature space. A systematic leakage audit computes each candidate feature's "
    "mutual information with Y, restricted to records where the feature is observable strictly before "
    "the funding decision; any feature whose informativeness collapses under the strict-temporal "
    "filter is dropped. The audit is documented in a feature catalogue committed to the project "
    "repository as the principal data-governance artefact, and accompanies a datasheet [Gebru et "
    "al., 2021] describing collection, provenance, intended uses, and known limitations including "
    "the originator's 'green' labelling taxonomy."
)

add_heading(doc, "5.4 Exploratory Data Analysis and Distribution-Shift Diagnostics", 2)
add_para(
    doc,
    "EDA proceeds in three layers. Univariate analysis produces distributional summaries (mean, "
    "median, standard deviation, skewness, kurtosis), Kolmogorov–Smirnov tests against fitted "
    "parametric families, and Shannon entropy for categoricals. Bivariate analysis computes Pearson "
    "and Spearman correlations, mutual information, and conditional default rates segmented by "
    "Industry, Borrower State, and Current Tier — identifying segments with anomalous default "
    "frequencies. Multivariate analysis uses PCA and UMAP projections for visual cluster inspection, "
    "and variance-inflation-factor (VIF) screening for multicollinearity. Distribution-shift "
    "diagnostics stratify the data temporally (Q1–Q4 of 2019) and compute the Population Stability "
    "Index (PSI) on each feature across quarters, detecting drift that could compromise model "
    "deployment beyond the training window. Class-conditional histograms identify features whose "
    "distributional separability suggests predictive value; per-feature ROC provides a one-feature "
    "predictive ceiling for sanity-checking later model contributions."
)

add_heading(doc, "5.5 Preprocessing Pipeline", 2)
add_para(
    doc,
    "The preprocessing pipeline is built as a scikit-learn ColumnTransformer to guarantee identical "
    "transformation at train and inference time and to prevent train–test leakage. Stage 1 — "
    "missing-data treatment: features with more than 40% missingness are dropped; for the remainder, "
    "median imputation is applied to numerics, an explicit 'missing' category is added for "
    "categoricals, and missing-indicator binary features are appended where missingness is itself "
    "informative (for example, missing Time in Business correlates with newer businesses). Stage 2 — "
    "outlier handling: winsorisation at the 1st/99th percentile for unbounded financial features, "
    "with sensitivity tests at 0.5/99.5 and no-winsorisation control runs. Stage 3 — encoding: "
    "one-hot encoding for low-cardinality categoricals (≤10 levels), target encoding with "
    "leave-one-out cross-fitting [Micci-Barreca, 2001] for high-cardinality categoricals (Industry, "
    "Borrower State), and frequency encoding for ordinal-like fields. Stage 4 — scaling: "
    "StandardScaler for distance-based learners (logistic regression, SVM, MLP) and identity "
    "(no scaling) for tree-based learners, which are scale-invariant. Stage 5 — feature interactions: "
    "domain-motivated derived features including loan-to-revenue ratio, payback-to-monthly-sales "
    "ratio, factor-rate-adjusted APR, and time-to-decision in days."
)

add_heading(doc, "5.6 Feature Engineering and Selection", 2)
add_para(
    doc,
    "A two-stage selection protocol is applied. Stage 1 is a filter: mutual-information ranking "
    "against Y, with the bottom decile dropped. Stage 2 is a wrapper: Boruta with shadow-feature "
    "null distribution [Kursa and Rudnicki, 2010] using a Random Forest oracle, run in parallel with "
    "a SHAP-importance ranking on a tuned XGBoost; features are required to clear both thresholds. "
    "The intersection produces a parsimonious feature set (target: 15–25 features) that maximises "
    "performance subject to interpretability. Stability of the selected set is verified by repeating "
    "Stage 2 across 30 bootstrap resamples and reporting the selection-frequency distribution; "
    "features selected in fewer than 60% of bootstraps are dropped as unstable. The original "
    "12-feature shortlist used in the supervisor draft is included as a benchmark to confirm or "
    "supersede."
)

add_heading(doc, "5.7 Class Imbalance Strategy", 2)
add_para(
    doc,
    "Three balancing strategies are compared head-to-head on a held-out validation fold: (i) "
    "class-weighted loss (no resampling; scale_pos_weight for XGBoost, class_weight='balanced' for "
    "sklearn learners); (ii) SMOTE-NC [Chawla et al., 2002] for mixed numeric/categorical data, "
    "applied strictly inside each training fold during cross-validation; (iii) focal loss [Lin et "
    "al., 2017] for the neural baselines. The strategy minimising the joint PR-AUC × calibration-"
    "error objective on validation is adopted per model family. Resampling is never applied to "
    "validation or test folds."
)

add_heading(doc, "5.8 Model Architectures", 2)
add_para(doc, "Six classifiers are trained, organised into three families for controlled comparison.", italic=True)

add_para(doc, "Baseline (linear, regulatory-default class):", bold=True)
add_bullet(doc, "L1- and L2-regularised Logistic Regression with monotonic priors on directionally constrained features.")
add_bullet(doc, "Kernel SVM (RBF), included as a non-linear shallow learner — historically dominant in early credit-scoring ML literature [Baesens et al., 2003] and a useful sanity check.")

add_para(doc, "Tree ensembles (expected state of the art):", bold=True)
add_bullet(doc, "Random Forest [Breiman, 2001] as a variance-reduction ensemble baseline.")
add_bullet(doc, "XGBoost [Chen and Guestrin, 2016] as the primary candidate, with monotonic constraints on Credit Score, Annual Revenue, and Time in Business.")
add_bullet(doc, "LightGBM [Ke et al., 2017] and CatBoost [Prokhorenkova et al., 2018] as alternative GBDT implementations whose differing bias profiles serve as cross-validation against single-implementation idiosyncrasies.")

add_para(doc, "Tabular deep learning:", bold=True)
add_bullet(doc, "Multi-Layer Perceptron with batch normalisation and dropout, trained with focal loss.")
add_bullet(doc, "FT-Transformer [Gorishniy et al., 2021] with feature-tokenisation and 4 attention blocks — the strongest single tabular DL architecture in recent independent benchmarks.")

add_para(
    doc,
    "All models share an identical preprocessing pipeline (Section 5.5) so that differences reflect "
    "inductive bias rather than leakage or transformation drift between conditions."
)

add_heading(doc, "5.9 Training Protocol: Nested Cross-Validation and Bayesian Optimisation", 2)
add_para(
    doc,
    "Hyperparameter selection and performance estimation use 5×10 nested stratified k-fold "
    "cross-validation. The outer 10-fold loop estimates generalisation performance; the inner 5-fold "
    "loop performs hyperparameter optimisation via Optuna's Tree-structured Parzen Estimator [Akiba "
    "et al., 2019] with 100 trials per model per outer fold and HyperBand pruning to early-stop "
    "unpromising trials. Stratification is on the binary target. Fold splits are seeded and "
    "persisted, and the identical splits are used across all six model families to ensure paired "
    "comparisons. Statistical comparison of model AUCs uses DeLong's test [DeLong et al., 1988]; for "
    "calibration error and PR-AUC, paired bootstrap with 10,000 resamples produces 95% confidence "
    "intervals and one-sided p-values. The previous draft's flat 10-fold protocol is preserved as a "
    "benchmark and explicitly compared with the nested protocol to demonstrate the difference."
)

add_heading(doc, "5.10 Probability Calibration and Uncertainty Quantification", 2)
add_para(
    doc,
    "Even well-discriminating models are routinely mis-calibrated. For each candidate model, "
    "calibration is assessed pre-intervention via reliability diagrams, Brier score, and Expected "
    "Calibration Error (ECE) with 15 equal-frequency bins. Three calibration techniques are then "
    "compared on a held-out calibration fold: Platt scaling [Platt, 1999]; isotonic regression "
    "[Zadrozny and Elkan, 2002]; and temperature scaling for the neural learners. The selected "
    "calibrator is fit on a dedicated calibration split disjoint from both training and test. For "
    "deployment-time uncertainty, split-conformal prediction [Vovk et al., 2005; Angelopoulos and "
    "Bates, 2023] generates per-applicant prediction intervals with guaranteed marginal coverage at "
    "user-specified confidence levels (90%, 95%), without distributional assumptions — a property "
    "increasingly discussed in regulatory commentary on high-risk AI as a robustness primitive."
)

add_heading(doc, "5.11 Explainability Framework", 2)
add_para(doc, "EMERALD-AI integrates a three-layer explainability stack: global, local, and counterfactual.", italic=True)
add_para(
    doc,
    "Global. TreeSHAP mean-|φ| feature importance, SHAP interaction values to quantify second-order "
    "feature dependencies, partial-dependence (PDP) and accumulated-local-effects (ALE) plots for "
    "non-linear marginal responses, and monotonicity verification on the constrained features."
)
add_para(
    doc,
    "Local. Per-applicant TreeSHAP attributions surfaced as a waterfall chart in the web application, "
    "with a KernelSHAP cross-check on a 200-applicant sample to validate that TreeSHAP attributions "
    "do not diverge under feature-dependence assumptions [Aas et al., 2021], and LIME [Ribeiro et "
    "al., 2016] as a perturbation-based secondary explanation."
)
add_para(
    doc,
    "Counterfactual. DiCE [Mothilal et al., 2020] generates the minimal feature changes that would "
    "flip a declined applicant to approved — directly addressing the actionability requirement under "
    "GDPR Article 22 and the FCA Consumer Duty's 'support' outcome. Counterfactuals are constrained "
    "to actionable features (excluding, e.g., Borrower State) and tested for plausibility against the "
    "training distribution."
)
add_para(
    doc,
    "Explanation fidelity is empirically validated using the Quantus toolkit [Hedström et al., 2023], "
    "reporting faithfulness, robustness, and complexity scores for each attribution method — moving "
    "beyond the implicit assumption that an explanation is reliable simply because it is generated. "
    "All explainability artefacts ship with a Method Card [Adkins et al., 2022] declaring prescribed "
    "and proscribed uses, so downstream consumers know not only what the model does but what it "
    "should not be asked to do."
)

add_heading(doc, "5.12 Fairness, Stability, and Robustness Audit", 2)
add_para(
    doc,
    "Although the dataset does not contain protected attributes directly, indirect disparate impact "
    "is testable on three proxy axes: Industry, Borrower State, and business-size segments. For each "
    "axis, four fairness metrics are computed using the AIF360 toolkit [Bellamy et al., 2018]: (i) "
    "demographic-parity gap; (ii) equalised-odds gap (TPR and FPR differences across groups); (iii) "
    "predictive-parity gap (precision differences); (iv) calibration-within-group via group-"
    "stratified reliability diagrams. Disparity exceeding pre-registered thresholds triggers "
    "re-training with reweighting [Kamiran and Calders, 2012] or adversarial debiasing [Zhang et "
    "al., 2018]. The audit is conducted with Selbst et al.'s (2019) five abstraction traps as "
    "explicit guardrails: the chosen parity criterion is reported as a value judgement (calibration-"
    "within-group is selected as the binding constraint, given the regulatory weight of PD reporting "
    "under the Internal Ratings-Based approach), the base-rate decomposition is published so "
    "reviewers can disagree on visible grounds, and the fairness claim is explicitly conditioned on "
    "the deployment context rather than presented as portable. Robustness is assessed via three "
    "protocols: (a) input perturbation — Gaussian noise injection at varying SNR for numerics and "
    "single-feature category flips for categoricals, reporting prediction-flip rate; (b) "
    "leave-one-segment-out validation — training on all industries except one and testing on the "
    "held-out industry, estimating transfer; (c) temporal generalisation — training on H1 2019 and "
    "testing on H2 2019 with PSI-tracked drift diagnostics."
)

add_heading(doc, "5.13 Evaluation Metrics", 2)
add_para(
    doc,
    "The primary metric is PR-AUC, which is robust to class imbalance and aligns with credit "
    "decisioning's emphasis on minority-class recall. Secondary metrics: ROC-AUC, Kolmogorov–Smirnov "
    "statistic (industry-standard in credit risk), F1 at the operating point, recall@top-decile "
    "(regulator-relevant for adverse-action volume estimation), Brier score, Expected Calibration "
    "Error, and Matthews Correlation Coefficient. Operating-threshold selection optimises expected "
    "loss under a configurable cost matrix encoding the relative costs of false positives (lost good "
    "loans) and false negatives (defaults). All results are reported with bootstrap 95% confidence "
    "intervals and paired statistical tests across models on identical folds."
)

add_heading(doc, "5.14 Web Application and MLOps Architecture", 2)
add_para(
    doc,
    "EMERALD-AI is operationalised as a containerised microservice stack. The architecture is "
    "deliberately conservative: Paleyes et al. (2022) catalogue, on a survey of production-ML "
    "deployments, that the dominant failure modes are not modelling errors but data-pipeline, "
    "monitoring, and team-handoff issues — and Nahar et al. (2022) document the cross-role friction "
    "that emerges when ML, software, and product teams share an artefact. EMERALD-AI therefore "
    "treats the model-card, datasheet, and method-card handoff as a first-class deliverable rather "
    "than a documentation afterthought."
)
add_table(
    doc,
    headers=["Layer", "Technology", "Function"],
    rows=[
        ["Model Registry", "MLflow", "Versioning of trained models, metrics, artefacts, and preprocessing transformers"],
        ["Backend API", "Python • FastAPI • Uvicorn • Pydantic", "REST endpoints: /score, /batch_score, /explain, /counterfactual, /portfolio, /fairness_audit, /healthz"],
        ["ML Runtime", "scikit-learn • XGBoost • LightGBM • SHAP • DiCE • joblib", "Inference, SHAP attribution, counterfactual generation, conformal interval computation"],
        ["Frontend SPA", "React 18 • Vite • TypeScript • TanStack Query • Recharts", "Four pages: Portfolio Dashboard, Single Predict, Batch Score, SHAP Explorer"],
        ["Monitoring", "Prometheus • Grafana • Evidently", "Input-feature and prediction-distribution drift monitoring with PSI thresholds and alerting"],
        ["Containerisation", "Docker • docker-compose", "Reproducible local and CI deployment"],
        ["CI/CD", "GitHub Actions", "Lint, type-check, unit and integration tests, automated model-card regeneration on each merge"],
    ],
)
add_para(
    doc,
    "The Single Predict page renders a structured input form, returns the EMERALD-AI score with its "
    "conformal interval, a SHAP waterfall, and one or more counterfactuals ('what would change this "
    "decision?'). The Batch Score page accepts CSV uploads and returns scored output with per-row "
    "top-3 SHAP features. The Portfolio Dashboard surfaces real-time aggregated insights with a "
    "fairness-audit panel. The SHAP Explorer is the public model-card surface, exposing global "
    "feature importance, SHAP interaction matrices, and PDP/ALE plots. All endpoints emit structured "
    "logs to enable the post-market monitoring required by EU AI Act Article 61."
)

add_heading(doc, "5.15 Analytical Tools", 2)
add_table(
    doc,
    headers=["Tool / Library", "Purpose in EMERALD-AI"],
    rows=[
        ["Python 3.11", "Core language for all data processing, modelling, and backend development"],
        ["Pandas • NumPy • Polars", "Data ingestion, cleaning, transformation, and descriptive statistics"],
        ["scikit-learn • imbalanced-learn", "Linear baselines, RF, cross-validation, SMOTE-NC, metrics, calibrators"],
        ["XGBoost • LightGBM • CatBoost", "Gradient-boosted decision tree implementations with monotonic-constraint support"],
        ["PyTorch", "MLP and FT-Transformer training"],
        ["Optuna", "Bayesian hyperparameter optimisation with HyperBand pruning"],
        ["SHAP • DiCE • LIME • Quantus", "Explainability stack and explanation-fidelity validation"],
        ["AIF360", "Fairness metrics and bias-mitigation algorithms"],
        ["MAPIE", "Conformal prediction (split-conformal intervals)"],
        ["MLflow • DVC", "Experiment tracking, model registry, and data/artefact versioning"],
        ["FastAPI • Uvicorn", "Production-grade REST API layer"],
        ["React 18 • Vite • TypeScript • Recharts", "Frontend SPA with typed components and interactive visualisation"],
        ["Docker • GitHub Actions", "Containerisation and CI/CD"],
        ["Jupyter / Google Colab Pro+", "Exploratory analysis and GPU-accelerated DL training"],
    ],
)

add_heading(doc, "5.16 Reproducibility and Computational Infrastructure", 2)
add_para(
    doc,
    "All experimentation runs in a containerised Python 3.11 environment with pinned dependencies via "
    "a uv lockfile; data, model, and metric artefacts are versioned with DVC; experiment runs are "
    "tracked in MLflow with full hyperparameter, code-version, and seed lineage. Compute: "
    "hyperparameter search runs on Google Colab Pro+ A100 GPU instances for the DL models and "
    "Warwick HPC CPU nodes for the GBDT models, with budgets pre-registered (approximately 500 "
    "GPU-hours and 2,000 CPU-hours in total). A `make reproduce` target re-runs the full pipeline "
    "from raw data to scored test set in under eight hours of wall-clock time, satisfying both "
    "academic reproducibility norms [Pineau et al., 2021] and the EU AI Act's technical "
    "documentation requirements."
)

add_heading(doc, "5.17 Ethical Considerations", 2)
add_bullet(doc, "Data privacy. The dataset is de-identified by the providing institution prior to release; no personally identifying information enters analysis. All artefacts and processed extracts are stored in encrypted university storage and destroyed at submission, per Warwick's research ethics protocol and UK GDPR.")
add_bullet(doc, "Fairness. The audit framework of Section 5.12 is mandatory, not optional; failure to clear pre-registered fairness thresholds will be reported transparently in the dissertation alongside successful results.")
add_bullet(doc, "Human oversight. The web application explicitly badges all outputs as decision-support, not decision-replacement, and surfaces counterfactuals to enable substantive applicant recourse.")
add_bullet(doc, "Transparency. Code, model cards [Mitchell et al., 2019], method cards [Adkins et al., 2022], datasheets [Gebru et al., 2021], and the full training pipeline are released under a permissive open-source licence on submission, subject to the data licence's commercial-use restrictions.")

doc.add_page_break()

# ============================================================
# 8. EXPECTED CONTRIBUTIONS
# ============================================================
add_heading(doc, "6. Expected Contributions", 1)
add_para(doc, "Methodological", bold=True)
add_numbered(doc, "A reproducible end-to-end ML pipeline for green-loan credit scoring integrating nested CV, Bayesian HPO, calibration, conformal uncertainty, multi-method explainability, fairness audit, and an explicit reject-inference sensitivity protocol — currently absent in the literature.")
add_numbered(doc, "An empirical benchmark of six classifier families (LR, SVM, RF, XGBoost/LightGBM/CatBoost, FT-Transformer/MLP) on the same green-loan dataset under identical preprocessing — the first apples-to-apples comparison of tabular DL versus GBDT in the green-credit context.")

add_para(doc, "Substantive", bold=True)
add_numbered(doc, "Identification of the feature determinants of green-loan creditworthiness, including derivative features (loan-to-revenue ratio, factor-rate-adjusted APR) and segment-conditional risk heterogeneity that uniform underwriting thresholds currently miss.")
add_numbered(doc, "Empirical evidence on whether the climate–credit channel documented at the macro level [Mirza et al., 2024; Kölbel et al., 2022] is recoverable in transaction-level green-loan default outcomes, and whether modern tabular deep learning closes its historical gap to GBDTs at this data scale.")

add_para(doc, "Practical and Translational", bold=True)
add_numbered(doc, "The EMERALD-AI web application as an open-source, MLOps-grade reference implementation that lending institutions and researchers can adapt — bridging the production gap that has limited academic-to-industry knowledge transfer in credit AI [Paleyes et al., 2022].")
add_numbered(doc, "A regulator-ready audit template covering calibration, fairness, robustness, and explanation fidelity, mapped explicitly to EU AI Act Annex III obligations and FCA Consumer Duty outcomes, accompanied by model card, method card, and datasheet artefacts.")

# ============================================================
# 9. PROJECT PLAN
# ============================================================
add_heading(doc, "7. Project Plan (16 Weeks)", 1)
add_table(
    doc,
    headers=["Weeks", "Workstream", "Key Deliverables"],
    rows=[
        ["1–2", "Setup", "Final literature scan, data-access confirmation, environment & CI bootstrap, ethics approval"],
        ["3–4", "Data understanding", "EDA report, target-leakage audit, reject-inference sensitivity protocol, feature catalogue, datasheet"],
        ["5–7", "Data preparation", "Preprocessing pipeline, feature engineering, class-imbalance experiments"],
        ["8–10", "Modelling", "Nested CV and Bayesian HPO across all six model families; statistical comparison"],
        ["11", "Calibration & uncertainty", "Reliability diagrams, calibrator selection, conformal intervals"],
        ["12–13", "Audit", "Explainability stack with Quantus fidelity validation; fairness, stability, and robustness audit; method card"],
        ["14–15", "Deployment", "Web application implementation, MLOps integration, user-acceptance evaluation"],
        ["16", "Submission", "Dissertation write-up, model-card and datasheet finalisation, open-source release"],
    ],
)

# ============================================================
# 8. FUTURE WORK (briefly) — moved into the same section as references
# ============================================================
add_heading(doc, "8. Scope Boundaries and Future Work", 1)
add_para(
    doc,
    "EMERALD-AI's scope is bounded deliberately. Several adjacent research directions are each a "
    "thesis in their own right and are noted here only to make the boundary explicit. Federated and "
    "decentralised credit scoring across institutions, differential-privacy guarantees on per-"
    "applicant attributions, secure multi-party computation for cross-lender risk pooling, "
    "survival-analytic time-to-default modelling, and online / continual learning under regulatory-"
    "grade post-market monitoring would all be valuable extensions; none is in the dissertation's "
    "scope. The reject-inference treatment is also bounded: the dataset contains only funded loans, "
    "so the full toolkit [Banasik and Crook, 2007; Kang et al., 2021; Shen et al., 2022] cannot be "
    "exercised end-to-end and is documented as a stipulated requirement of any production "
    "deployment rather than a deliverable of this work."
)

# ============================================================
# 9. REFERENCES
# ============================================================
doc.add_page_break()
add_heading(doc, "9. References", 1)
add_para(doc, "Citations follow author–year style; the bibliography below covers v0.3's literature base (74 entries). New additions over v0.2 are flagged with the year tag of the bot-discovered batch (2026-05-18) — they are the climate-credit channel, the sociotechnical fairness critique, the reject-inference treatment, the production-ML deployment evidence, and the algorithmic-opacity framing.", italic=True)

refs = [
    "Aas, K., Jullum, M., & Løland, A. (2021). Explaining individual predictions when features are dependent. Artificial Intelligence, 298, 103502.",
    "Adkins, D., Alsallakh, B., Cheema, A., Kokhlikyan, N., McReynolds, E., Mishra, P., Procope, C., Sawruk, J., Wang, E., & Zvyagina, P. (2022). Method cards for prescriptive machine-learning transparency. ACM EAAMO.",
    "Akiba, T., Sano, S., Yanase, T., Ohta, T., & Koyama, M. (2019). Optuna: A next-generation hyperparameter optimization framework. KDD.",
    "Altman, E. I. (1968). Financial ratios, discriminant analysis and the prediction of corporate bankruptcy. Journal of Finance, 23(4), 589–609.",
    "Angelopoulos, A. N., & Bates, S. (2023). Conformal prediction: A gentle introduction. Foundations and Trends in ML, 16(4).",
    "Arik, S. Ö., & Pfister, T. (2021). TabNet: Attentive interpretable tabular learning. AAAI.",
    "Baesens, B., Van Gestel, T., Viaene, S., Stepanova, M., Suykens, J., & Vanthienen, J. (2003). Benchmarking state-of-the-art classification algorithms for credit scoring. Journal of the Operational Research Society, 54(6), 627–635.",
    "Banasik, J., & Crook, J. (2007). Reject inference, augmentation, and sample selection. European Journal of Operational Research, 183(3), 1582–1594.",
    "Beaver, W. H. (1966). Financial ratios as predictors of failure. Journal of Accounting Research, 4, 71–111.",
    "Bellamy, R. K. E., et al. (2018). AI Fairness 360: An extensible toolkit for detecting, understanding, and mitigating unwanted algorithmic bias. IBM Journal of R&D, 63(4/5).",
    "Berg, F., Kölbel, J. F., & Rigobon, R. (2022). Aggregate confusion: The divergence of ESG ratings. Review of Finance, 26(6), 1315–1344.",
    "Borisov, V., Leemann, T., et al. (2022). Deep neural networks and tabular data: A survey. IEEE Transactions on Neural Networks and Learning Systems.",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
    "Burrell, J. (2016). How the machine 'thinks': Understanding opacity in machine learning algorithms. Big Data & Society, 3(1), 1–12.",
    "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. JAIR, 16, 321–357.",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. KDD.",
    "Chouldechova, A. (2017). Fair prediction with disparate impact: A study of bias in recidivism prediction instruments. Big Data, 5(2), 153–163.",
    "Climate Bonds Initiative. (2023). Climate Bonds Standard v4.",
    "Dafermos, Y., & Nikolaidi, M. (2021). How can green differentiated capital requirements affect climate risks? A dynamic macrofinancial analysis. Journal of Financial Stability, 54, 100871.",
    "DeLong, E. R., DeLong, D. M., & Clarke-Pearson, D. L. (1988). Comparing the areas under two or more correlated ROC curves. Biometrics, 44(3), 837–845.",
    "Dumitrescu, E., Hue, S., Hurlin, C., & Tokpavi, S. (2022). Machine learning for credit scoring: Improving logistic regression with non-linear decision-tree effects. EJOR, 297(3), 1178–1192.",
    "Ehlers, T., & Packer, F. (2017). Green bond finance and certification. BIS Quarterly Review, September 2017.",
    "European Banking Authority. (2020). Guidelines on loan origination and monitoring (EBA/GL/2020/06).",
    "European Commission. (2020). Regulation (EU) 2020/852 on the establishment of a framework to facilitate sustainable investment (Taxonomy Regulation).",
    "European Commission. (2021). Proposal for a Regulation laying down harmonised rules on artificial intelligence (Artificial Intelligence Act). COM/2021/206 final.",
    "Financial Conduct Authority. (2022). Consumer Duty: Final rules and guidance. Policy Statement PS22/9.",
    "Flammer, C. (2021). Corporate green bonds. Journal of Financial Economics, 142(2), 499–516.",
    "Gan, X. D., Zheng, X. Y., Li, C. C., & Zhu, G. Q. (2023). Green bond issuance and trade credit access: Evidence from Chinese bond market. Finance Research Letters, 58, 104842.",
    "Gebru, T., Morgenstern, J., et al. (2021). Datasheets for datasets. Communications of the ACM, 64(12), 86–92.",
    "Gorishniy, Y., Rubachev, I., Khrulkov, V., & Babenko, A. (2021). Revisiting deep learning models for tabular data. NeurIPS.",
    "Grinsztajn, L., Oyallon, E., & Varoquaux, G. (2022). Why do tree-based models still outperform deep learning on tabular data? NeurIPS Datasets and Benchmarks.",
    "Hand, D. J., & Henley, W. E. (1997). Statistical classification methods in consumer credit scoring: A review. Journal of the Royal Statistical Society A, 160(3), 523–541.",
    "Hardt, M., Price, E., & Srebro, N. (2016). Equality of opportunity in supervised learning. NeurIPS.",
    "He, H., Bai, Y., Garcia, E. A., & Li, S. (2008). ADASYN: Adaptive synthetic sampling approach for imbalanced learning. IJCNN.",
    "Hedström, A., et al. (2023). Quantus: An explainable AI toolkit for responsible evaluation of neural network explanations. JMLR, 24.",
    "Henley, W. E. (1995). Statistical aspects of credit scoring. PhD thesis, The Open University.",
    "Janzing, D., Minorics, L., & Blöbaum, P. (2020). Feature relevance quantification in explainable AI: A causal problem. AISTATS.",
    "Kamiran, F., & Calders, T. (2012). Data preprocessing techniques for classification without discrimination. Knowledge and Information Systems, 33(1), 1–33.",
    "Kang, Y., Jia, N., Cui, R., & Deng, J. (2021). A graph-based semi-supervised reject inference framework considering imbalanced data distribution for consumer credit scoring. Applied Soft Computing, 105, 107259.",
    "Ke, G., Meng, Q., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. NeurIPS.",
    "Kleinberg, J., Mullainathan, S., & Raghavan, M. (2017). Inherent trade-offs in the fair determination of risk scores. ITCS.",
    "Kölbel, J. F., Leippold, M., Rillaerts, J., & Wang, Q. (2022). Ask BERT: How regulatory disclosure of transition and physical climate risks affects the CDS term structure. Journal of Financial Econometrics, 22(1), 30–69.",
    "Kursa, M. B., & Rudnicki, W. R. (2010). Feature selection with the Boruta package. Journal of Statistical Software, 36(11).",
    "Lessmann, S., Baesens, B., Seow, H. V., & Thomas, L. C. (2015). Benchmarking state-of-the-art classification algorithms for credit scoring: An update of research. EJOR, 247(1), 124–136.",
    "Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2017). Focal loss for dense object detection. ICCV.",
    "Loan Market Association, APLMA, & LSTA. (2023). Green Loan Principles.",
    "Lu, J., Liu, A., Dong, F., Gu, F., Gama, J., & Zhang, G. (2018). Learning under concept drift: A review. IEEE TKDE.",
    "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. NeurIPS.",
    "Lundberg, S. M., et al. (2020). From local explanations to global understanding with explainable AI for trees. Nature Machine Intelligence, 2, 56–67.",
    "Micci-Barreca, D. (2001). A preprocessing scheme for high-cardinality categorical attributes in classification and prediction problems. SIGKDD Explorations, 3(1), 27–32.",
    "Mirza, N., Umar, M., Horobeţ, A., & Boubaker, S. (2024). Effects of climate change and technological capex on credit risk cycles in the European Union. Technological Forecasting and Social Change, 204, 123448.",
    "Mitchell, M., et al. (2019). Model cards for model reporting. FAT* 2019, 220–229.",
    "Mothilal, R. K., Sharma, A., & Tan, C. (2020). Explaining machine learning classifiers through diverse counterfactual explanations. FAccT.",
    "Nahar, N., Zhou, S., Lewis, G. A., & Kästner, C. (2022). Collaboration challenges in building ML-enabled systems. ICSE 2022, 413–425.",
    "Paleyes, A., Urma, R.-G., & Lawrence, N. D. (2022). Challenges in deploying machine learning: A survey of case studies. ACM Computing Surveys, 55(6), 114.",
    "Pasquale, F. (2015). The Black Box Society: The Secret Algorithms That Control Money and Information. Harvard University Press.",
    "Pineau, J., et al. (2021). Improving reproducibility in machine learning research. JMLR, 22.",
    "Platt, J. (1999). Probabilistic outputs for support vector machines and comparisons to regularized likelihood methods. Advances in Large Margin Classifiers.",
    "Popov, S., Morozov, S., & Babenko, A. (2020). Neural oblivious decision ensembles for deep learning on tabular data. ICLR.",
    "Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). CatBoost: Unbiased boosting with categorical features. NeurIPS.",
    "Rawal, K., & Lakkaraju, H. (2020). Beyond individualized recourse: Interpretable and interactive summaries of actionable recourses. NeurIPS.",
    "Reboredo, J. C., & Ugolini, A. (2020). Price connectedness between green bond and financial markets. Economic Modelling, 88, 25–38.",
    "Ribeiro, M. T., Singh, S., & Guestrin, C. (2016). 'Why should I trust you?': Explaining the predictions of any classifier. KDD.",
    "Rudin, C. (2019). Stop explaining black box machine learning models for high stakes decisions and use interpretable models instead. Nature Machine Intelligence, 1, 206–215.",
    "Santos, M. S., Soares, J. P., Abreu, P. H., Araújo, H., & Santos, J. (2018). Cross-validation for imbalanced datasets: Avoiding overoptimistic and overfitting approaches. IEEE Computational Intelligence Magazine, 13(4).",
    "Selbst, A. D., boyd, d., Friedler, S. A., Venkatasubramanian, S., & Vertesi, J. (2019). Fairness and abstraction in sociotechnical systems. FAT* 2019, 59–68.",
    "Shen, F., Yang, Z., Zhao, X., & Lan, D. (2022). Reject inference in credit scoring using a three-way decision and safe semi-supervised support vector machine. Information Sciences, 606, 614–627.",
    "Shwartz-Ziv, R., & Armon, A. (2022). Tabular data: Deep learning is not all you need. Information Fusion, 81.",
    "Slack, D., Hilgard, S., Jia, E., Singh, S., & Lakkaraju, H. (2020). Fooling LIME and SHAP: Adversarial attacks on post hoc explanation methods. AIES.",
    "Somepalli, G., Goldblum, M., Schwarzschild, A., Bruss, C. B., & Goldstein, T. (2021). SAINT: Improved neural networks for tabular data via row attention and contrastive pre-training. arXiv:2106.01342.",
    "Tang, D. Y., & Zhang, Y. (2020). Do shareholders benefit from green bonds? Journal of Corporate Finance, 61, 101427.",
    "Thomas, L. C. (2000). A survey of credit and behavioural scoring. International Journal of Forecasting, 16(2), 149–172.",
    "Vovk, V., Gammerman, A., & Shafer, G. (2005). Algorithmic learning in a random world. Springer.",
    "Wachter, S., Mittelstadt, B., & Russell, C. (2017). Counterfactual explanations without opening the black box: Automated decisions and the GDPR. Harvard Journal of Law & Technology, 31(2).",
    "Wirth, R., & Hipp, J. (2000). CRISP-DM: Towards a standard process model for data mining. PAKDD.",
    "Zadrozny, B., & Elkan, C. (2002). Transforming classifier scores into accurate multiclass probability estimates. KDD.",
    "Zhang, B. H., Lemoine, B., & Mitchell, M. (2018). Mitigating unwanted biases with adversarial learning. AIES.",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

doc.save("proposal_third_draft.docx")
print("OK: proposal_third_draft.docx")
