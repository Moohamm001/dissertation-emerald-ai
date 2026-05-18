"""Build the 3K-word EMERALD-AI dissertation proposal (v0.4.1 condensed-3k).

Tightest companion to `build_proposal.py`. Target: ~3,000 words total
(body ~2,200 + curated 30-ref bibliography). Same v0.4 facts and v0.4.1
conformal reframe; references reduced to the 30 essential entries cited
in the condensed body.

Preserves the v0.4 surgical edits unchanged:
  - §4.4 empirical 0.36% delinquent prevalence (vs textbook 2-15%)
  - §5.2 label counts (3,848/10,124/49/1/113 = 99.20% coverage) and
    paidOff-only headline / all-favourable sensitivity reframing
  - §5.3 90 permitted / 76 forbidden audit decomposition + MCCC
    100%-missing transparency note + datasheet cross-reference
  - §5.5 >40%-missing drop-list casualties (Term 86.4% / APR 59.6%
    / Factor 42.0%) + the 11 100%-missing columns
  - v0.4.1 §5.6 conformal reframe: marginal coverage as headline,
    Mondrian / class-conditional as diagnostic, interval width
    excluded from primary metrics, transparency-mechanism framing.

Output: proposal_fourth_draft_3k.docx.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x2E)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# =====================================================================
# TITLE BLOCK
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(
    "EMERALD-AI: An Explainable, Calibrated, and Audit-Ready\n"
    "Machine Learning Framework for Green Loan Credit Scoring,\n"
    "Operationalised as a Full-Stack Decision-Support Platform"
)
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x14, 0x4F, 0x32)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("MSc Applied Artificial Intelligence — Proposal (v0.4.1, 3K condensed)")
sr.italic = True
sr.font.size = Pt(11)

doc.add_paragraph()

# =====================================================================
# ABSTRACT
# =====================================================================
add_heading(doc, "Abstract", 1)
add_para(
    doc,
    "Green-loan origination has grown at >40% CAGR since 2018, yet credit-risk infrastructure "
    "for sustainable lending uses scorecards designed for conventional consumer credit. "
    "Recent evidence shows climate and clean-tech capex now alter credit-risk cycles in "
    "asset-class-specific ways [Mirza et al., 2024; Kölbel et al., 2022]. EMERALD-AI proposes "
    "an explainable, calibrated, audit-ready ML framework trained on a real 2019 dataset of "
    "14,135 funded green-loan transactions (166 columns; 90 features retained after the "
    "leakage audit). It benchmarks six classifier families under nested CV with Bayesian HPO, "
    "integrates post-hoc calibration and split-conformal uncertainty re-scoped to the small-"
    "minority regime (50 defaults; 0.36% prevalence), delivers a three-layer explainability "
    "stack with Quantus fidelity validation [Hedström et al., 2023], conducts a fairness "
    "audit on industry/state/size proxies, and ships as an open-source FastAPI + React "
    "platform aligned with the EU AI Act and FCA Consumer Duty."
)

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
add_heading(doc, "1. Introduction", 1)
add_para(
    doc,
    "Green loans ring-fence proceeds for environmentally beneficial purposes — renewables, "
    "retrofits, clean transport, sustainable agriculture — and the market has grown >40% "
    "CAGR since 2018. Credit assessment still relies on conventional scorecards built on a "
    "narrow set of indicators, poorly equipped for back-loaded renewable cash flows or "
    "greenwashing exposure. The mis-pricing is empirically observable: climate and clean-tech "
    "capex alter EU credit-risk cycles [Mirza et al., 2024]; climate-risk disclosures move CDS "
    "term structure [Kölbel et al., 2022]. EMERALD-AI is an explainable, calibrated, audit-"
    "ready ML credit-scoring framework trained on 14,135 funded 2019 green-loan transactions "
    "(166 structured variables), delivered alongside a production web application so the "
    "research is directly usable by lending institutions."
)

# =====================================================================
# 2. MOTIVATION
# =====================================================================
add_heading(doc, "2. Motivation", 1)
add_para(
    doc,
    "Three pressures converge. Operationally, lenders re-purpose conventional scorecards on "
    "green portfolios without modelling green-asset cash-flow specificity or the heterogeneity "
    "of ESG taxonomies [Berg, Kölbel & Rigobon, 2022]. Academically, no published study on "
    "real green-loan data simultaneously delivers modern tabular benchmarking, post-hoc "
    "calibration with conformal uncertainty, multi-method XAI with empirical fidelity "
    "validation, sociotechnically-disciplined fairness audit on green-lending proxies, "
    "explicit reject-inference treatment, and a deployable lending-officer interface. "
    "Regulatorily, the EU AI Act places credit scoring under Annex III high-risk obligations "
    "(risk management, data governance, transparency, human oversight, accuracy logging, "
    "post-market monitoring), FCA Consumer Duty makes outcome-based explainability binding on "
    "UK lenders, and EBA/GL/2020/06 mandates sector-specific robustness and explainability. "
    "Better-calibrated green-credit decisions also improve capital allocation to genuine "
    "climate solutions [Flammer, 2021]."
)

# =====================================================================
# 3. RESEARCH OBJECTIVES AND QUESTIONS
# =====================================================================
add_heading(doc, "3. Research Objectives and Research Questions", 1)
add_para(doc, "Objectives:", bold=True)
add_bullet(doc, "RO1. Benchmark six classifier families (regularised LR, RBF-SVM, RF, XGBoost, LightGBM, CatBoost, MLP, FT-Transformer) on real green-loan data under identical preprocessing.")
add_bullet(doc, "RO2. Validate post-hoc calibration (Platt/isotonic/temperature) and split-conformal uncertainty against nominal coverage, re-scoped to the small-minority regime.")
add_bullet(doc, "RO3. Build a three-layer explainability stack (TreeSHAP global+local, LIME, DiCE counterfactuals) and validate via Quantus fidelity metrics.")
add_bullet(doc, "RO4. Conduct a fairness/robustness audit on industry/state/business-size proxies aligned with EU AI Act Annex III and FCA Consumer Duty.")
add_bullet(doc, "RO5. Deploy as an open-source FastAPI + React + MLflow platform.")
add_para(doc, "Questions:", bold=True)
add_bullet(doc, "RQ1. Which classifier family wins under identical preprocessing on green-loan data, measured on PR-AUC, within-minority-class ECE, and recall@top-decile?")
add_bullet(doc, "RQ2. Does post-hoc calibration close within-minority ECE gaps and does split-conformal coverage meet its nominal marginal target at 0.36% prevalence?")
add_bullet(doc, "RQ3. Do TreeSHAP, LIME, and DiCE agree on the top contributing features per decision, and which delivers highest Quantus fidelity?")
add_bullet(doc, "RQ4. Do disparities in demographic parity, equalised odds, predictive parity, and calibration-within-group persist across industry/state/size after audit-stage mitigations?")

# =====================================================================
# 4. LITERATURE REVIEW
# =====================================================================
add_heading(doc, "4. Literature Review", 1)
add_para(
    doc,
    "Credit scoring evolved from discriminant analysis [Altman, 1968] through logistic "
    "regression to statistical-learning benchmarks [Baesens et al., 2003; Lessmann et al., "
    "2015; Dumitrescu et al., 2022]. Gradient-boosted trees are the contemporary tabular "
    "state of the art [Chen & Guestrin, 2016; Ke et al., 2017; Prokhorenkova et al., 2018], "
    "retaining the lead over tabular deep learning on most benchmarks [Grinsztajn et al., "
    "2022], though FT-Transformer [Gorishniy et al., 2021] is competitive on larger datasets."
)
add_para(
    doc,
    "Class imbalance, calibration, uncertainty, and selection bias dominate the credit-risk "
    "methodology stack. SMOTE-family resampling [Chawla et al., 2002] must be confined within "
    "CV folds [Santos et al., 2018]; cost-sensitive learning and focal loss [Lin et al., "
    "2017] are alternatives. Post-hoc calibration and split-conformal prediction "
    "[Angelopoulos & Bates, 2023] are largely absent from green-credit literature. Accepted-"
    "only selection bias [Banasik & Crook, 2007] has been modernised by graph-based semi-"
    "supervised propagation [Kang et al., 2021], which presupposes access to rejected "
    "applicants' features."
)
add_para(
    doc,
    "Green-finance specificity matters: ESG rating divergence [Berg, Kölbel & Rigobon, 2022] "
    "makes definitional inputs heterogeneous; green-bond pricing [Flammer, 2021] supplies the "
    "macro context. Explainability in high-stakes finance spans intrinsically interpretable "
    "models [Rudin, 2019], TreeSHAP [Lundberg et al., 2020], and counterfactuals [Mothilal, "
    "Sharma & Tan, 2020]; fidelity must be validated, not assumed [Hedström et al., 2023]."
)
add_para(
    doc,
    "Fairness in classification is bound by inherent trade-offs [Kleinberg, Mullainathan & "
    "Raghavan, 2017; Chouldechova, 2017]; mitigation toolkits exist [Bellamy et al., 2018]; "
    "sociotechnical critique cautions against treating fairness as a pure abstraction [Selbst "
    "et al., 2019]. Documentation primitives include model cards [Mitchell et al., 2019] and "
    "datasheets [Gebru et al., 2021]."
)
add_para(
    doc,
    "Research gap. No published study on real green-loan data simultaneously delivers all "
    "six of: modern tabular benchmarking under identical preprocessing; post-hoc calibration "
    "+ conformal uncertainty; multi-method XAI with empirical fidelity validation; "
    "sociotechnically-disciplined fairness audit on green-lending proxies; explicit reject-"
    "inference treatment of accepted-only selection bias; and a deployable lending-officer "
    "interface. EMERALD-AI occupies precisely this gap."
)

# =====================================================================
# 5. METHODOLOGY
# =====================================================================
add_heading(doc, "5. Methodology", 1)

add_heading(doc, "5.1 Problem Formalisation and Label Construction", 2)
add_para(
    doc,
    "Let X ∈ ℝ^d be the feature vector and Y ∈ {0,1} the binary creditworthiness indicator, "
    "Y=1 favourable (paidOff ∪ current), Y=0 delinquent (default ∪ behind). Deal Status "
    "supervises 14,022 of 14,135 records (99.20% coverage): 3,848 paidOff (27.4%) + 10,124 "
    "current (72.2%) on the favourable side, 49 default (0.35%) + 1 behind (0.01%) delinquent, "
    "113 unlabelled. The 0.36% delinquent prevalence is a substantive empirical constraint "
    "(§5.4). Three label-construction risks are addressed:"
)
add_bullet(doc, "Censoring bias. Current loans are right-censored — their observation window pre-dates loan-term completion. paidOff-only is the headline robustness check; all-favourable is the optimistic comparator; both numbers are reported.")
add_bullet(doc, "Definition leakage. Post-funding fields are excluded from X (§5.2 audit).")
add_bullet(doc, "Accepted-only selection bias. Banasik & Crook (2007) showed this materially biases parameter estimates; the modern toolkit [Kang et al., 2021] requires rejected-applicant features that the dataset lacks. The model is framed as conditional on the prior accept-policy rather than as a portable underwriting function.")

add_heading(doc, "5.2 Dataset Characterisation and Target-Leakage Audit", 2)
add_para(
    doc,
    "The 2019 All Funded Green Loan dataset is 14,135 transactions × 166 columns. The leakage "
    "audit (v1.0; `data/governance/datasheet.md`, `feature_catalogue.yaml`, "
    "`feature_audit_summary.md`) classifies every column into one of six categories: 23 pre-"
    "funding applicant, 15 pre-funding loan-offer, 9 structural metadata, 43 deal-progression "
    "timestamp, 28 post-funding outcome, and 48 administrative / free-text — yielding 90 "
    "permitted features and 76 forbidden. Load-bearing exclusions include Deal Status "
    "(defines Y), Term Complete Percentage, Percent Paid, the Closed family, and the Is "
    "Offer Received/Accepted/Published flags. Transparency note: `Monthly Credit Card "
    "Charges`, treated as load-bearing in earlier scoping, is 100% missing. Three further "
    "permitted features exceed the >40%-missing drop threshold: Term (86.4%, also semantically "
    "overloaded), APR (59.6%), Factor (42.0%). Eleven columns are wholly missing (App Out, "
    "1st Online Engmnt, Used Online Experience, Closed Lenders, MCCC, Rep Type, Rep Is Active, "
    "Inactive Status, Closed By Type, Dead Status, Renewal Eligible Date). The dissertation "
    "surfaces these rather than silently dropping them."
)

add_heading(doc, "5.3 EDA, Preprocessing, Feature Selection", 2)
add_para(
    doc,
    "EDA proceeds in three layers: univariate (distributional summaries, KS tests), bivariate "
    "(Pearson/Spearman, mutual information, conditional default rates by Industry / Borrower "
    "State / Current Tier), and multivariate (PCA, UMAP, VIF). Temporal stratification (Q1–"
    "Q4 2019) plus Population Stability Index detects drift. The preprocessing pipeline is "
    "a ColumnTransformer: Stage 1 drops features >40% missing (removing Term, APR, Factor + "
    "the eleven 100%-missing columns), median-imputes numerics, adds 'missing' categories, "
    "appends missing-indicators where informative; Stage 2 winsorises at 1/99th percentile "
    "with sensitivity controls; Stage 3 one-hot encodes low-cardinality categoricals, target-"
    "encodes high-cardinality; Stage 4 scales for distance-based learners only; Stage 5 "
    "derives domain-motivated interactions. Selection is a two-stage MI-filter + Boruta / "
    "SHAP-importance wrapper with 30-bootstrap stability filtering."
)

add_heading(doc, "5.4 Class Imbalance Strategy", 2)
add_para(
    doc,
    "The empirical 0.36% delinquent prevalence sits an order of magnitude below the typical "
    "2–15% range — only 50 minority observations exist. SMOTE-family oversampling [Chawla "
    "et al., 2002] interpolates between extremely sparse seeds and risks off-manifold "
    "synthetic points; it is included for completeness but compared head-to-head against "
    "cost-sensitive learning (scale_pos_weight, class_weight='balanced'), focal-loss training "
    "[Lin et al., 2017], and a one-class / anomaly-detection baseline. All resampling occurs "
    "strictly within CV folds [Santos et al., 2018]. Raw accuracy is uninformative (a constant "
    "predictor scores 99.64%) and is explicitly excluded from the metric panel."
)

add_heading(doc, "5.5 Models, Training, Calibration, Uncertainty", 2)
add_para(
    doc,
    "Six classifier families are trained: regularised logistic regression and RBF-SVM "
    "(regulatory-default baselines); Random Forest [Breiman, 2001], XGBoost [Chen & "
    "Guestrin, 2016] (primary candidate, with monotonic constraints on Credit Score, Annual "
    "Revenue, Time in Business), LightGBM [Ke et al., 2017], CatBoost [Prokhorenkova et al., "
    "2018] (tree ensembles); MLP and FT-Transformer [Gorishniy et al., 2021] (tabular deep "
    "learning). Training uses 5×10 nested stratified CV with Bayesian HPO. Calibration is "
    "selected on a held-out split (Platt, isotonic, temperature) via reliability diagrams; "
    "within-minority-class ECE is the headline calibration metric — marginal ECE is dominated "
    "by the favourable class and masks the failures that matter for adverse-action "
    "decisioning."
)
add_para(
    doc,
    "Split-conformal prediction [Angelopoulos & Bates, 2023] supplies marginal coverage at "
    "90% and 95%, with the finite-sample exact guarantee P(Y ∈ C(X)) = ⌈(N_cal+1)(1−α)⌉ / "
    "(N_cal+1) well-supported by a ~2,800-row calibration split. Class-conditional / "
    "Mondrian conformal coverage is reported as a diagnostic with explicit small-sample "
    "bootstrap CIs; interval width is excluded from primary metrics because, with ~10 "
    "minority observations in the calibration set, it reflects measurement-precision rather "
    "than model quality. Conformal is positioned as a regulator-facing transparency "
    "mechanism — the intervals communicate honest small-sample uncertainty — rather than a "
    "precision claim the data cannot defend."
)

add_heading(doc, "5.6 Explainability, Fairness, Robustness Audit", 2)
add_para(
    doc,
    "Three layers: global (TreeSHAP [Lundberg et al., 2020]), local (KernelSHAP + LIME), "
    "and counterfactual (DiCE [Mothilal, Sharma & Tan, 2020]). Fidelity, sensitivity, "
    "robustness, and complexity are evaluated with Quantus [Hedström et al., 2023]. Fairness "
    "audit uses AIF360 [Bellamy et al., 2018] across Industry, Borrower State, and a derived "
    "business-size proxy: demographic parity, equalised odds, predictive parity, calibration-"
    "within-group. The Selbst et al. (2019) sociotechnical traps frame the audit: the parity "
    "criterion is reported as a value judgement (calibration-within-group as the binding "
    "constraint), base-rate decomposition is published, and fairness claims are conditioned "
    "on the deployment context rather than presented as portable. Stability uses 30-bootstrap "
    "selection-frequency and AUC variance; robustness uses adversarial perturbations, leave-"
    "one-segment-out validation, and H1/H2-2019 temporal generalisation."
)

add_heading(doc, "5.7 Evaluation, Deployment, Reproducibility, Ethics", 2)
add_para(
    doc,
    "Primary metrics (each reported with 10,000-bootstrap 95% CIs): PR-AUC, within-minority-"
    "class ECE, recall@top-decile. Secondary (cross-study comparability only): ROC-AUC, KS, "
    "F1, marginal Brier, marginal ECE, conformal marginal coverage, MCC. Conformal class-"
    "conditional coverage and interval width are diagnostics. Deployment is FastAPI + "
    "React/Vite SPA, MLflow for runs, Prometheus + Evidently for drift monitoring — "
    "providing the post-market monitoring layer EU AI Act Article 61 anticipates. "
    "Reproducibility: uv-locked deps, DVC-tracked data, seeded RNGs, model-card + datasheet "
    "[Gebru et al., 2021] + method-card [Mitchell et al., 2019] regenerated on each merge. "
    "Ethics: Warwick REC protocol; data de-identified; no direct protected attributes — "
    "fairness audit uses indirect proxies disclosed in §5.6."
)

# =====================================================================
# 6. EXPECTED CONTRIBUTIONS
# =====================================================================
add_heading(doc, "6. Expected Contributions", 1)
add_bullet(doc, "Methodological: first integrated benchmark on real green-loan data combining identical-preprocessing comparison across six classifier families with calibration, conformal uncertainty (small-N-honest framing), multi-method XAI + Quantus fidelity, sociotechnical fairness audit, and explicit reject-inference framing.")
add_bullet(doc, "Substantive: empirical characterisation of the 2019 green-loan portfolio under extreme imbalance (0.36% delinquent), the leakage-audit-validated 90-feature space, and the documented Term/APR/Factor data-quality limitations.")
add_bullet(doc, "Practical: open-source MLOps-grade reference implementation (EMERALD-AI web app) plus a regulator-ready audit template mapped to EU AI Act Annex III and FCA Consumer Duty.")

# =====================================================================
# 7. PROJECT PLAN
# =====================================================================
add_heading(doc, "7. Project Plan (16 Weeks)", 1)
add_table(
    doc,
    headers=["Weeks", "Workstream", "Key Deliverables"],
    rows=[
        ["1–2", "Setup", "Literature scan, data access, environment + CI, ethics approval"],
        ["3–4", "Data understanding", "EDA report, leakage audit, datasheet, feature catalogue"],
        ["5–7", "Data preparation", "Preprocessing pipeline, feature engineering, imbalance experiments"],
        ["8–10", "Modelling", "Nested CV + Bayesian HPO; statistical comparison"],
        ["11", "Calibration & uncertainty", "Reliability diagrams, calibrator selection, conformal intervals"],
        ["12–13", "Audit", "Explainability + Quantus fidelity; fairness/stability/robustness"],
        ["14–15", "Deployment", "Web application, MLOps integration, user-acceptance evaluation"],
        ["16", "Submission", "Dissertation write-up; documentation finalisation; open-source release"],
    ],
)

# =====================================================================
# 8. SCOPE BOUNDARIES AND FUTURE WORK
# =====================================================================
add_heading(doc, "8. Scope Boundaries and Future Work", 1)
add_para(
    doc,
    "Out-of-scope but recorded as future work: federated and decentralised credit scoring, "
    "differential-privacy guarantees on attributions, secure multi-party computation for "
    "cross-lender pooling, survival-analytic time-to-default modelling, and online / "
    "continual learning under regulatory-grade post-market monitoring. The reject-inference "
    "treatment is bounded by data availability: only funded loans are present, so the full "
    "toolkit [Banasik & Crook, 2007; Kang et al., 2021] cannot be exercised end-to-end."
)

# =====================================================================
# 9. REFERENCES (30 essential — those cited in this condensed body)
# =====================================================================
doc.add_page_break()
add_heading(doc, "9. References", 1)
add_para(doc, "Curated 30-entry bibliography for the 3K-word condensed version. The full 74-ref bibliography is retained in the v0.4.1 supervisor-facing draft (proposal_fourth_draft.docx).", italic=True)

refs = [
    "Altman, E. I. (1968). Financial ratios, discriminant analysis and the prediction of corporate bankruptcy. Journal of Finance, 23(4), 589–609.",
    "Angelopoulos, A. N., & Bates, S. (2023). Conformal prediction: A gentle introduction. Foundations and Trends in ML, 16(4).",
    "Baesens, B., Van Gestel, T., Viaene, S., Stepanova, M., Suykens, J., & Vanthienen, J. (2003). Benchmarking state-of-the-art classification algorithms for credit scoring. Journal of the Operational Research Society, 54(6), 627–635.",
    "Banasik, J., & Crook, J. (2007). Reject inference, augmentation, and sample selection. European Journal of Operational Research, 183(3), 1582–1594.",
    "Bellamy, R. K. E., et al. (2018). AI Fairness 360: An extensible toolkit for detecting, understanding, and mitigating unwanted algorithmic bias. IBM Journal of R&D, 63(4/5).",
    "Berg, F., Kölbel, J. F., & Rigobon, R. (2022). Aggregate confusion: The divergence of ESG ratings. Review of Finance, 26(6), 1315–1344.",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
    "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. JAIR, 16, 321–357.",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. KDD.",
    "Chouldechova, A. (2017). Fair prediction with disparate impact: A study of bias in recidivism prediction instruments. Big Data, 5(2), 153–163.",
    "Dumitrescu, E., Hue, S., Hurlin, C., & Tokpavi, S. (2022). Machine learning for credit scoring: Improving logistic regression with non-linear decision-tree effects. European Journal of Operational Research, 297(3), 1178–1192.",
    "Flammer, C. (2021). Corporate green bonds. Journal of Financial Economics, 142(2), 499–516.",
    "Gebru, T., Morgenstern, J., et al. (2021). Datasheets for datasets. Communications of the ACM, 64(12), 86–92.",
    "Gorishniy, Y., Rubachev, I., Khrulkov, V., & Babenko, A. (2021). Revisiting deep learning models for tabular data. NeurIPS.",
    "Grinsztajn, L., Oyallon, E., & Varoquaux, G. (2022). Why do tree-based models still outperform deep learning on tabular data? NeurIPS Datasets and Benchmarks.",
    "Hedström, A., et al. (2023). Quantus: An explainable AI toolkit for responsible evaluation of neural network explanations. Journal of Machine Learning Research, 24.",
    "Kang, Y., Jia, N., Cui, R., & Deng, J. (2021). A graph-based semi-supervised reject inference framework considering imbalanced data distribution for consumer credit scoring. Applied Soft Computing, 105, 107259.",
    "Ke, G., Meng, Q., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. NeurIPS.",
    "Kleinberg, J., Mullainathan, S., & Raghavan, M. (2017). Inherent trade-offs in the fair determination of risk scores. ITCS.",
    "Kölbel, J. F., Leippold, M., Rillaerts, J., & Wang, Q. (2022). Ask BERT: How regulatory disclosure of transition and physical climate risks affects the CDS term structure. Journal of Financial Econometrics, 22(1), 30–69.",
    "Lessmann, S., Baesens, B., Seow, H. V., & Thomas, L. C. (2015). Benchmarking state-of-the-art classification algorithms for credit scoring: An update of research. European Journal of Operational Research, 247(1), 124–136.",
    "Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2017). Focal loss for dense object detection. ICCV.",
    "Lundberg, S. M., et al. (2020). From local explanations to global understanding with explainable AI for trees. Nature Machine Intelligence, 2, 56–67.",
    "Mirza, N., Umar, M., Horobeţ, A., & Boubaker, S. (2024). Effects of climate change and technological capex on credit risk cycles in the European Union. Technological Forecasting and Social Change, 204, 123448.",
    "Mitchell, M., et al. (2019). Model cards for model reporting. FAT* 2019, 220–229.",
    "Mothilal, R. K., Sharma, A., & Tan, C. (2020). Explaining machine learning classifiers through diverse counterfactual explanations. FAccT.",
    "Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). CatBoost: Unbiased boosting with categorical features. NeurIPS.",
    "Rudin, C. (2019). Stop explaining black box machine learning models for high stakes decisions and use interpretable models instead. Nature Machine Intelligence, 1, 206–215.",
    "Santos, M. S., Soares, J. P., Abreu, P. H., Araújo, H., & Santos, J. (2018). Cross-validation for imbalanced datasets: Avoiding overoptimistic and overfitting approaches. IEEE Computational Intelligence Magazine, 13(4).",
    "Selbst, A. D., boyd, d., Friedler, S. A., Venkatasubramanian, S., & Vertesi, J. (2019). Fairness and abstraction in sociotechnical systems. FAT* 2019, 59–68.",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

doc.save("proposal_fourth_draft_3k.docx")
print("OK: proposal_fourth_draft_3k.docx")
